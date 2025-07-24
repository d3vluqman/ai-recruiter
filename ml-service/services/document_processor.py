import io
import logging
from typing import Optional
from fastapi import UploadFile
import PyPDF2
import pdfplumber
from docx import Document
import re

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for extracting text from various document formats"""

    def __init__(self):
        self.supported_formats = [".pdf", ".doc", ".docx", ".txt"]

    async def extract_text_from_file(self, file: UploadFile) -> str:
        """
        Extract text content from uploaded file

        Args:
            file: UploadFile object from FastAPI

        Returns:
            str: Extracted text content

        Raises:
            ValueError: If file format is not supported
            Exception: If text extraction fails
        """
        try:
            # Read file content
            content = await file.read()
            file_extension = self._get_file_extension(file.filename)

            # Reset file pointer for potential re-reading
            await file.seek(0)

            if file_extension == ".pdf":
                return self._extract_from_pdf(content)
            elif file_extension in [".doc", ".docx"]:
                return self._extract_from_docx(content)
            elif file_extension == ".txt":
                return self._extract_from_txt(content)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")

        except Exception as e:
            logger.error(f"Failed to extract text from {file.filename}: {str(e)}")
            raise Exception(f"Text extraction failed: {str(e)}")

    def _get_file_extension(self, filename: str) -> str:
        """Get file extension from filename"""
        if not filename:
            raise ValueError("Filename is required")
        return "." + filename.lower().split(".")[-1] if "." in filename else ""

    def _extract_from_pdf(self, content: bytes) -> str:
        """
        Extract text from PDF using multiple methods for better accuracy

        Args:
            content: PDF file content as bytes

        Returns:
            str: Extracted text
        """
        text = ""

        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            # If pdfplumber didn't extract much text, try PyPDF2
            if len(text.strip()) < 100:
                text = ""
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

        except Exception as e:
            logger.warning(f"PDF extraction error: {str(e)}")
            # Fallback to PyPDF2 if pdfplumber fails
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as fallback_error:
                logger.error(
                    f"Both PDF extraction methods failed: {str(fallback_error)}"
                )
                raise Exception("Failed to extract text from PDF")

        return self._clean_text(text)

    def _extract_from_docx(self, content: bytes) -> str:
        """
        Extract text from DOCX file

        Args:
            content: DOCX file content as bytes

        Returns:
            str: Extracted text
        """
        try:
            doc = Document(io.BytesIO(content))
            text = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

            return self._clean_text(text)

        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise Exception("Failed to extract text from DOCX file")

    def _extract_from_txt(self, content: bytes) -> str:
        """
        Extract text from TXT file

        Args:
            content: TXT file content as bytes

        Returns:
            str: Extracted text
        """
        try:
            # Try different encodings
            encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]

            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return self._clean_text(text)
                except UnicodeDecodeError:
                    continue

            # If all encodings fail, use utf-8 with error handling
            text = content.decode("utf-8", errors="ignore")
            return self._clean_text(text)

        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            raise Exception("Failed to extract text from TXT file")

    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text

        Args:
            text: Raw extracted text

        Returns:
            str: Cleaned text
        """
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r"\s+", " ", text)

        # Remove special characters that might interfere with parsing
        text = re.sub(r"[^\w\s@.,;:()\-+/&%$#]", " ", text)

        # Normalize line breaks
        text = re.sub(r"\n+", "\n", text)

        # Remove leading/trailing whitespace
        text = text.strip()

        return text

    def validate_extracted_text(self, text: str, min_length: int = 50) -> bool:
        """
        Validate that extracted text meets minimum quality requirements

        Args:
            text: Extracted text to validate
            min_length: Minimum required text length

        Returns:
            bool: True if text is valid, False otherwise
        """
        if not text or len(text.strip()) < min_length:
            return False

        # Check if text contains meaningful content (not just special characters)
        meaningful_chars = re.sub(r"[^\w\s]", "", text)
        if len(meaningful_chars.strip()) < min_length * 0.7:
            return False

        return True
